import docx

d = docx.Document()
d.add_heading("Python Test File")
d.add_paragraph("Hello dont sleep IMP Topic")
d.add_picture(r"C:\Users\Naveen\Downloads\sample_images\seven.jpg")
d.save("dummy.docx")
